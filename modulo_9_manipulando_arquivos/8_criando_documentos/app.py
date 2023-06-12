from docx import Document
from docx.shared import Cm

documento = Document()# Cria uma instância de um documento

documento.add_heading('Título do documento', 0) # Cria o título do documento

paragrafo = documento.add_paragraph('Um parágrafo simples ') # Cria um parágrafo
# Concatenação de textos no parágrafo
paragrafo.add_run('e importante').bold = True
paragrafo.add_run(' com palavras em ')
paragrafo.add_run('itálico').italic = True

# Cria títulos por nível
documento.add_heading('Título - Nível 1', level=1)
documento.add_heading('Título - Nível 2', level=2)
documento.add_heading('Título - Nível 3', level=3)

# Adiciona estilo nos parágrafos
documento.add_paragraph('Formatação "No Spacing"', style='No Spacing')
documento.add_paragraph('Formatação "Heading 1"', style='Heading 1')
documento.add_paragraph('Formatação "Heading 2"', style='Heading 2')
documento.add_paragraph('Formatação "Heading 3"', style='Heading 3')
documento.add_paragraph('Formatação "Title"', style='Title')
documento.add_paragraph('Formatação "Subtitle"', style='Subtitle')
documento.add_paragraph('Formatação "Quote"', style='Quote')
documento.add_paragraph('Formatação "Intense Quote"', style='Intense Quote')
documento.add_paragraph('Formatação "List Paragraph"', style='List Paragraph')
documento.add_paragraph('Formatação "List Bullet"', style='List Bullet')
documento.add_paragraph('Formatação "List Number"', style='List Number')

# Adicionando imagem no documento
documento.add_picture('notebook.jpg', width=Cm(5.25))

# Adicionando tabela no documento
tabela = documento.add_table(rows=3, cols=2)
# Adiciona valores na tabela
celula1 = tabela.cell(0,0)
celula1.text = 'Nome'

# Outra forma de criar uma tabela no documento com várias informações
compras = (
    (3, '101', 'Uva'),
    (7, '422', 'Pão'),
    (4, '423', 'Banana, Ovos, Carne, Tomate')
)

tabela_supermercado = documento.add_table(rows=1, cols=3)
cabecalho_tabela_supermercado = tabela_supermercado.rows[0].cells
cabecalho_tabela_supermercado[0].text = 'Quantidade'
cabecalho_tabela_supermercado[1].text = 'Id'
cabecalho_tabela_supermercado[2].text = 'Descrição'

for quantidade, id, desc in compras:
    linha = tabela_supermercado.add_row().cells
    linha[0].text = str(quantidade)
    linha[1].text = id
    linha[2].text = desc

# criar quebra de página no documento
documento.add_page_break()

documento.add_paragraph('Estamos em outra página')

documento.save('demo.docx') # Salva o documento